from __future__ import annotations

import base64
import re
from pathlib import Path

from pidraw.detector import detect
from pidraw.renderer import render

_FENCED_BLOCK = re.compile(
    r"```(\w+)?\n(.*?)```",
    re.DOTALL,
)

_DIAGRAM_LANGUAGES = {
    "mermaid",
    "mmd",
    "plantuml",
    "puml",
    "iuml",
    "graphviz",
    "dot",
    "gv",
    "d2",
    "bpmn",
    "nomnoml",
    "wavedrom",
    "vega",
    "vega-lite",
}


def extract_diagrams(markdown: str) -> list[dict]:
    blocks: list[dict] = []
    for match in _FENCED_BLOCK.finditer(markdown):
        lang = (match.group(1) or "").strip().lower()
        code = match.group(2).strip()
        if not code:
            continue
        is_diagram = lang in _DIAGRAM_LANGUAGES
        if not is_diagram and lang in ("", "text", "plain"):
            detected = detect(code)
            is_diagram = detected != "unknown"
            if is_diagram:
                lang = detected
        if is_diagram:
            blocks.append(
                {
                    "match": match,
                    "language": lang,
                    "code": code,
                    "start": match.start(),
                    "end": match.end(),
                }
            )
    return blocks


def render_diagrams(
    blocks: list[dict],
    fmt: str = "svg",
    scale: float = 1.0,
    transparent: bool = True,
) -> list[dict]:
    for block in blocks:
        try:
            result = render(
                block["code"],
                language=block["language"],
                format=fmt,
                scale=scale,
                transparent=transparent,
            )
            block["rendered"] = result
        except Exception as exc:
            block["error"] = str(exc)
    return blocks


def to_html(
    markdown: str,
    blocks: list[dict],
    fmt: str = "svg",
    title: str = "PiDraw Document",
) -> str:
    parts: list[str] = []
    last_end = 0

    for block in blocks:
        parts.append(markdown[last_end : block["start"]])
        parts.append(_render_block_html(block, fmt))
        last_end = block["end"]

    parts.append(markdown[last_end:])

    body = "".join(parts)
    return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>{_escape_html(title)}</title>
<style>
  body {{ font-family: system-ui, sans-serif; max-width: 960px; margin: 0 auto; padding: 20px; line-height: 1.6; }}
  img {{ max-width: 100%; height: auto; display: block; margin: 1em 0; border: 1px solid #e0e0e0; border-radius: 4px; }}
  svg {{ max-width: 100%; height: auto; display: block; margin: 1em 0; }}
  pre {{ background: #f5f5f5; padding: 1em; border-radius: 6px; overflow-x: auto; }}
  code {{ font-family: 'SF Mono', 'Fira Code', monospace; font-size: 0.9em; }}
</style>
</head>
<body>
{body}
</body>
</html>"""


def _render_block_html(block: dict, fmt: str) -> str:
    if "error" in block:
        return f"<pre>Diagram error: {_escape_html(block['error'])}</pre>\n"

    rendered = block["rendered"]
    if fmt == "png":
        b64 = base64.b64encode(rendered).decode("ascii")
        return f'<img src="data:image/png;base64,{b64}" alt="diagram">\n'
    return rendered + "\n"


def to_markdown(
    markdown: str,
    blocks: list[dict],
    fmt: str = "svg",
) -> str:
    parts: list[str] = []
    last_end = 0

    for block in blocks:
        parts.append(markdown[last_end : block["start"]])
        parts.append(_render_block_md(block, fmt))
        last_end = block["end"]

    parts.append(markdown[last_end:])
    return "".join(parts)


def _render_block_md(block: dict, fmt: str) -> str:
    if "error" in block:
        return f"```\nERROR: {block['error']}\n```\n"

    rendered = block["rendered"]
    if fmt == "png":
        b64 = base64.b64encode(rendered).decode("ascii")
        return f"![diagram](data:image/png;base64,{b64})\n\n"
    return f"```svg\n{rendered}\n```\n"


def _escape_html(text: str) -> str:
    return (
        text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")
    )


def _add_image_to_paragraph(paragraph, stream) -> None:
    from docx.shared import Inches
    from PIL import Image

    Image.MAX_IMAGE_PIXELS = None

    img = Image.open(stream)
    orig_w, orig_h = img.size
    stream.seek(0)

    page_w = 6.0
    max_h = 7.5

    if orig_w / orig_h > page_w / max_h:
        paragraph.add_run().add_picture(stream, width=Inches(page_w))
    else:
        paragraph.add_run().add_picture(stream, height=Inches(max_h))

    paragraph.paragraph_format.space_after = Inches(0.15)


def to_docx(
    markdown: str,
    blocks: list[dict],
    fmt: str = "png",
    title: str = "PiDraw Document",
) -> bytes:
    from io import BytesIO

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    doc = Document()

    _configure_styles(doc)

    last_end = 0

    for block in blocks:
        text_before = markdown[last_end : block["start"]]
        if text_before.strip():
            _add_text_to_docx(doc, text_before)

        if "error" in block:
            p = doc.add_paragraph()
            run = p.add_run(f"[Diagram error: {block['error']}]")
            run.font.size = Pt(8)
        elif fmt == "png":
            img_bytes = block["rendered"]
            if isinstance(img_bytes, bytes):
                from PIL import Image as PILImage

                PILImage.MAX_IMAGE_PIXELS = None
                preview = PILImage.open(BytesIO(img_bytes))
                blank = min(preview.size) < 20
                preview.close()
                if blank:
                    p = doc.add_paragraph()
                    run = p.add_run(f"[Diagram rendered blank: {block.get('language', '?')}]")
                    run.font.size = Pt(8)
                    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
                else:
                    stream = BytesIO(img_bytes)
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(4)
                    p.paragraph_format.space_after = Pt(4)
                    _add_image_to_paragraph(p, stream)

        last_end = block["end"]

    text_after = markdown[last_end:]
    if text_after.strip():
        _add_text_to_docx(doc, text_after)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _configure_styles(doc) -> None:
    from docx.shared import Pt

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(4)

    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = "Calibri"
        hs.paragraph_format.space_before = Pt(10)
        hs.paragraph_format.space_after = Pt(4)


def _add_text_to_docx(doc, text: str) -> None:
    from docx.shared import Pt

    paragraphs = re.split(r"\n{2,}", text)
    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            continue

        lines = para_text.splitlines()
        cleaned_lines = [l for l in lines if l.strip() and not l.strip().startswith("```")]
        if not cleaned_lines:
            continue
        combined = " ".join(l.strip() for l in cleaned_lines)

        if combined.startswith("### "):
            p = doc.add_heading(combined[4:], level=3)
        elif combined.startswith("## "):
            p = doc.add_heading(combined[3:], level=2)
        elif combined.startswith("# "):
            p = doc.add_heading(combined[2:], level=1)
        elif combined.startswith("- ") or combined.startswith("* "):
            for line in cleaned_lines:
                stripped = line.strip()
                if stripped.startswith("- ") or stripped.startswith("* "):
                    p = doc.add_paragraph(stripped[2:], style="List Bullet")
                elif stripped and not stripped.startswith("```"):
                    p = doc.add_paragraph(stripped)
                    p.paragraph_format.space_after = Pt(2)
        elif combined[0].isdigit() and ". " in combined[:4]:
            for line in cleaned_lines:
                stripped = line.strip()
                if stripped and not stripped.startswith("```"):
                    p = doc.add_paragraph(stripped, style="List Number")
        else:
            p = doc.add_paragraph(combined)
            p.paragraph_format.space_after = Pt(4)


def process_markdown(
    source: str,
    *,
    fmt: str = "svg",
    scale: float = 1.0,
    transparent: bool = True,
) -> tuple[str, list[dict]]:
    blocks = extract_diagrams(source)
    if not blocks:
        return source, []
    render_diagrams(blocks, fmt=fmt, scale=scale, transparent=transparent)
    return source, blocks


def process_markdown_file(
    path: str,
    *,
    fmt: str = "svg",
    scale: float = 1.0,
    transparent: bool = True,
) -> tuple[str, list[dict]]:
    source = Path(path).read_text(encoding="utf-8-sig")
    return process_markdown(source, fmt=fmt, scale=scale, transparent=transparent)


def render_md(
    source: str,
    *,
    output_format: str = "html",
    fmt: str = "svg",
    scale: float = 1.0,
    transparent: bool = True,
) -> str | bytes:
    markdown, blocks = process_markdown(source, fmt=fmt, scale=scale, transparent=transparent)
    if output_format == "html":
        return to_html(markdown, blocks, fmt=fmt)
    if output_format == "docx":
        return to_docx(markdown, blocks, fmt=fmt)
    return to_markdown(markdown, blocks, fmt=fmt)


def render_md_file(
    path: str,
    *,
    output_format: str = "html",
    fmt: str = "svg",
    scale: float = 1.0,
    transparent: bool = True,
) -> str | bytes:
    source = Path(path).read_text(encoding="utf-8-sig")
    return render_md(
        source,
        output_format=output_format,
        fmt=fmt,
        scale=scale,
        transparent=transparent,
    )
